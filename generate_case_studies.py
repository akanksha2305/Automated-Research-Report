# generate_case_studies.py

from docx import Document
import os

SAVE_DIR = "data/raw/case_studies/"
os.makedirs(SAVE_DIR, exist_ok=True)

sample_cases = [
    "A retail company adapts to AI-powered recommendations.",
    "A healthcare startup integrates a new diagnostics model.",
    "An EdTech firm scales using transformer-based chatbots.",
    # add more summaries...
]

def generate_case_studies(n=8):
    for i in range(n):
        doc = Document()
        doc.add_heading(f"Case Study {i+1}", 0)
        doc.add_paragraph(sample_cases[i % len(sample_cases)])
        doc.add_paragraph("\nDetails of the implementation, challenges, and outcomes...")
        doc.save(os.path.join(SAVE_DIR, f"case_study_{i+1}.docx"))

generate_case_studies()
