from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from docx import Document as DocxWriter
from content_synthesis import test_synthesis

# Template Design
def create_pdf_report(sections, filename="report.pdf"):
    pdf = canvas.Canvas(filename, pagesize=letter)
    width, height = letter
    pdf.drawString(100, height - 100, "Research Report")
    y = height - 150
    for section, content in sections.items():
        pdf.drawString(100, y, section)
        y -= 20
        pdf.drawString(100, y, content)
        y -= 100
    pdf.save()

def create_word_report(sections, filename="report.docx"):
    doc = DocxWriter()
    doc.add_heading('Research Report', 0)
    for section, content in sections.items():
        doc.add_heading(section, level=1)
        doc.add_paragraph(content)
    doc.save(filename)

# Document Automation
def generate_reports(query):
    sections = test_synthesis(query)
    create_pdf_report(sections)
    create_word_report(sections)

# Testing and QA
def test_report_generation():
    query = "AI research"
    generate_reports(query)

# Example usage
if __name__ == "__main__":
    test_report_generation()