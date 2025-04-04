import streamlit as st
from content_synthesis import test_synthesis
from docx import Document
from fpdf import FPDF

# === Page Setup ===
st.set_page_config(page_title="📘 Research Report Generator", layout="centered")
st.title("📘 Research Report Generator ✨")
st.markdown("<br>", unsafe_allow_html=True)
st.markdown("""
<div style='text-align: center; font-size:18px;'>
    🚀 Generate structured research reports powered by LLMs and RAG magic!<br>
    📄 Supports Word & PDF Export · 🔍 Semantic Understanding · 💡 Intelligent Insights<br>
</div>
""", unsafe_allow_html=True)
st.markdown("<br>", unsafe_allow_html=True)
st.markdown("📌 **This app supports queries related to:** 🧠 Deep Learning, 🤖 Artificial Intelligence, and 📈 Classic ML.")
st.markdown("<br>", unsafe_allow_html=True)
st.markdown("🧠 Enter your query below to generate a structured research report. You can download it as a Word or PDF document 📑.")
st.markdown("<br>", unsafe_allow_html=True)
# === Query Input ===
query = st.text_input("🔎 Enter research query")

# === Button Trigger ===
if st.button("🛠️ Generate Research Report") and query:
    sections = test_synthesis(query)
    st.success("✅ Report generated successfully!")
    st.subheader("📝 Report Preview")

    emoji_map = {
        "Abstract": "🧠",
        "Literature Review": "📚",
        "Methodology": "🛠️",
        "Findings": "🔍",
        "Conclusion": "✅"
    }

    for section, content in sections.items():
        emoji = emoji_map.get(section, "📄")
        st.markdown(f"### {emoji} {section}")
        st.write(content)

    # === Save Word ===
    def save_word():
        doc = Document()
        doc.add_heading('Research Report', 0)
        for section, content in sections.items():
            doc.add_heading(section, level=1)
            doc.add_paragraph(str(content))
        word_path = "research_report.docx"
        doc.save(word_path)
        with open(word_path, "rb") as f:
            st.download_button("⬇️ Download Word Report", f, file_name=word_path)

    # === Save PDF ===
    def save_pdf():
        from unicodedata import normalize
        def clean_text(text):
            return normalize("NFKD", text).encode("latin-1", "ignore").decode("latin-1")

        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.cell(200, 10, txt="Research Report", ln=True, align="C")
        for section, content in sections.items():
            pdf.ln(10)
            pdf.set_font("Arial", size=12, style='B')
            pdf.multi_cell(0, 10, clean_text(section))
            pdf.set_font("Arial", size=12)
            pdf.multi_cell(0, 10, clean_text(str(content)))
        pdf_path = "research_report.pdf"
        pdf.output(pdf_path)
        with open(pdf_path, "rb") as f:
            st.download_button("⬇️ Download PDF Report", f, file_name=pdf_path)

    save_word()
    save_pdf()


